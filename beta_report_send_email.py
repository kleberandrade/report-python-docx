# -*- coding: utf-8 -*-
"""
Created on Sat Nov 10 11:23:28 2018

@author: Kleber de Oliveira Andrade
@email: pdjkleber@gmail.com
"""

# -----------------------------------------------------------------------------
# BIBLIOTECAS PARA ENVIAR EMAIL
# -----------------------------------------------------------------------------
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.base import MIMEBase
from email.mime.text import MIMEText
from email.mime.image import MIMEImage
from email.utils import formatdate
from email import encoders

# -----------------------------------------------------------------------------
# BIBLIOTECAS UTILIZADAS PARA GRÁFICOS E LISTAS
# -----------------------------------------------------------------------------
import matplotlib.pyplot as plt
import pandas as pd
import numpy as np
import numpy.core.defchararray as np_f

# -----------------------------------------------------------------------------
# BIBLIOTECAS PARA CRIAÇÃO DO DOCX (PYTHON-DOCX)
# -----------------------------------------------------------------------------
from docx import Document
from docx.shared import Cm
from docx.shared import Pt
from docx.shared import RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# -----------------------------------------------------------------------------
# BIBLIOTECA PARA NÚVEM DE PALAVRAS
# -----------------------------------------------------------------------------
from wordcloud import WordCloud, STOPWORDS


# -----------------------------------------------------------------------------
# OUTRAS BIBLIOTECAS (POUCO UTILIZADAS NESTE PROJETO)
# -----------------------------------------------------------------------------
import os
from datetime import date
# limpa os arquivos da pasta conforme .bat
"""
    Informações do arquivo clear.bat
    del *.png
    del *.docx
"""
os.system('clear.bat')

# -----------------------------------------------------------------------------
# CONSTANTES UTILIZADAS PARA CRIAÇÃO DO RELATÓRIO (DOCX)
# -----------------------------------------------------------------------------
REPORT_HEADING = '3.2.\tRELATÓRIO DO JOGO VERSÃO BETA'

REPORT_PARAGRAPH_1 = 'Pesquisa realizada no dia 10 de novembro de 2018 com os alunos da Faculdade de Tecnologia de Americana na disciplina de \"Jogos para Console\", a fim de, avaliar a versão beta de cada jogo desenvolvido como proposta de atividade prática do semestre. O questionário apresenta 2 questões demográficas, 6 questões de múltipla escolha e 2 dissertativas.'
REPORT_PARAGRAPH_2 = 'O jogo \"{0}\", foi testado e avaliado por um grupo de {1} pessoas, composto por {2} mulheres e {3} homens com média de idade de {4} anos. A Figura 1 apresenta um gráfico de barras com a média e desvio padrão das seguintes avaliações: história ({5}), música e efeitos sonoros ({6}), jogabilidade ({7}), interface do usuário ({8}) e gráficos ({9}).'
REPORT_PARAGRAPH_3 = 'Para concluir a análise do questionário, foi identificada qual a dificuldade que o jogo apresentou para os participantes. A Figura 4 apresenta o percentual de dificuldade do jogo: {0}.'
REPORT_PARAGRAPH_4 = 'Segundo os participantes, o jogo se destaca pelos seguintes pontos fortes (resposta copiada na íntegra):'
REPORT_PARAGRAPH_5 = 'Para complementar a análise do jogo beta, os participantes ressaltaram os seguintes pontos fracos (resposta copiada na íntegra):'
REPORT_PARAGRAPH_6 = 'A Figura 2 apresenta uma nuvem de palavras resumindo os comentários dos pontos fortes do jogo avaliado.'
REPORT_PARAGRAPH_7 = 'A Figura 3 apresenta uma nuvem de palavras resumindo os comentários dos pontos fracos do jogo avaliado.'

REPORT_CAPTION_1 = 'Figura 1. Gráfico com as avaliações beta dos participantes.'
REPORT_CAPTION_2 = 'Figura 4. Gráfico do percentual de dificuldade do jogo por participantes.'
REPORT_CAPTION_3 = 'Figura 2. Resumo dos pontos fortes em formato de nuvem de palavras.' 
REPORT_CAPTION_4 = 'Figura 3. Resumo dos pontos fracos em formato de nuvem de palavras.' 

REPORT_OWN_QUOTATION = 'Fonte: autoria própria ({0}).'.format(int(date.today().year))


# -----------------------------------------------------------------------------
# FUNÇÕES UTILIZADAS PARA CRIAÇÃO DO RELATÓRIO (DOCX)
# -----------------------------------------------------------------------------

def add_document_heading(document, heading):
    
    # Definição de um estilo para a seção do capítulo
    style = document.styles['Body Text 2']  
    font = style.font
    font.name = 'Arial'
    font.size = Pt(12)
    font.all_caps = True
    font.bold = True
    font.color.rgb = RGBColor(0x0, 0x0, 0x0)
    
    # Adiciona uma seção no documento
    document.add_paragraph(REPORT_HEADING, 'Body Text 2')  
    
    
def add_document_paragraph(document, text):
    
    # Definição de um estilo para o parágrafo
    style = document.styles['Body Text']    
    font = style.font
    font.name = 'Arial'
    font.size = Pt(12)
    
    # Adiciona um paragrafo no documento
    paragraph = document.add_paragraph(text, 'Body Text')
    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.first_line_indent = Cm(1.25)


def add_document_figure(document, figure, caption, width):
    
    # Definição de um estilo para a legenda da figura
    style = document.styles['Caption']    
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    font.bold = False
    font.color.rgb = RGBColor(0x0, 0x0, 0x0)
    
    # Adiciona legenda para figura
    document.add_paragraph()
    document.add_paragraph(caption, 'Caption')    
    
    # Adiciona uma figura no documento
    document.add_picture(figure, width=Cm(width))
    last_paragraph = document.paragraphs[-1] 
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Adiciona a fonte da figura (citação)
    document.add_paragraph(REPORT_OWN_QUOTATION, 'Caption')
    document.add_paragraph()
    
    
def add_document_list(document, itens):
    
    # Definição de um estilo para a lista de itens
    style = document.styles['List Bullet']    
    font = style.font
    font.name = 'Arial'
    font.size = Pt(12)
    
    # Adiciona os itens (textos) no documento
    for text in list(set(itens)):
        paragraph = document.add_paragraph(str(text).rstrip('\r').replace('\n', '; '), 'List Bullet')
        paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.first_line_indent = Cm(1.25)
    

# -----------------------------------------------------------------------------
# CONSTANTES DAS COLUNAS DO ARQUIVO EXCEL
# -----------------------------------------------------------------------------
COLUMN_GAME = 'Jogo avaliado'
COLUMN_AGES = 'Idade'
COLUMN_GENDERS = 'Sexo'
COLUMN_GENDERS_ITENS = ['Masculino', 'Feminino']
COLUMN_GRAPHICS = 'Gráficos'
COLUMN_AUDIO = 'Música e SFX'
COLUMN_HUD = 'HUD'
COLUMN_GAMEPLAY = 'Jogabilidade'
COLUMN_STORY = 'História'
COLUMN_DIFFICULTY = 'Dificuldade'
COLUMN_DIFFICULTY_ITENS = ['Muito fácil', 'Fácil', 'Normal', 'Díficil', 'Muito Díficil']
COLUMN_STRONG = 'Ponto forte'
COLUMN_WEAK = 'Ponto fraco'

# -----------------------------------------------------------------------------
# FUNÇÕES UTILIZADAS PARA CRIAÇÃO DOS GRÁFICOS
# -----------------------------------------------------------------------------

def full_frame(width=None, height=None):
    import matplotlib as mpl
    mpl.rcParams['savefig.pad_inches'] = 0
    figsize = None if width is None else (width, height)
    fig = plt.figure(figsize=figsize)
    ax = plt.axes([0,0,1,1], frameon=False)
    ax.get_xaxis().set_visible(False)
    ax.get_yaxis().set_visible(False)
    #plt.autoscale(tight=True)
    return fig
    
def plot_bar_err(filename, data, legends):
    
    x_pos = [i for i, _ in enumerate(legends)]
    average = [None] * len(x_pos)
    variance = [None] * len(x_pos)
    
    for i in x_pos:
        average[i] = np.mean(data[i])
        variance[i] = np.std(data[i])
        
    print("[{}] Média: {} - ({})".format(filename, average, np.average(average)))
        
    plt.barh(x_pos, average, color='steelblue', xerr=variance)
    
    plt.yticks(x_pos, legends, fontsize=12)
    plt.xticks(list(range(11)), fontsize=12)
    
    plt.xlabel('Nota')
    
    plt.savefig(filename + '_bar_err.png', bbox_inches='tight', dpi=400)    
    plt.show()
    plt.close()
    
    
def convert_array_text_to_sequential_text(data):
    text = ''
    for i in range(len(data)):
        text = text + ' {0} '.format(data.replace('\\',''))
        if i < len(data) - 2:
            text = text + ' , '
        elif i < len(data) - 1:
            text = text + ' e '
            
    return text.lower()

    
def plot_word_cloud(filename, data):

    fig = full_frame()
    
    stopwords = set([line.rstrip() for line in open('stopwords.txt')])
    
    text = convert_array_text_to_sequential_text(data)
    wordcloud = WordCloud(stopwords=stopwords, width=1280, height=680, max_font_size=400, max_words=200, collocations=False, colormap='Blues').generate(text)
    
    plt.imshow(wordcloud, interpolation="bilinear")
    
    plt.savefig(filename,  bbox_inches='tight', pad_inches=0, dpi=900)   
    plt.show()
    plt.close()
    
    
def plot_pie_difficulty(filename, data):
    
    # Gráfico de pizza (direita) correspondente a quanto pagariam pelo jogo
    data = np.array(COLUMN_DIFFICULTY_ITENS)[data - 1]
    labels, counts_elements = np.unique(data, return_counts=True)
    slices = np.round(counts_elements / len(data) * 100,1 )
    
    # Convertendo os labels para string e adicionando caracter de escape no $ (latex)
    labels_str = labels.astype(str)
    labels_str = np_f.replace(labels_str, '$', '\$')
    
    patches, texts, autotexts = plt.pie(slices, labels=labels_str, autopct='%.1f%%', startangle=90, counterclock=False)

    plt.setp(autotexts, size=16, weight="bold") 
    plt.axis('equal')

    for t in texts:
        t.set_size('large')
    for t in autotexts:
        t.set_size('large')

    plt.savefig(filename, bbox_inches='tight', dpi=400)    
    plt.show()
    plt.close()
    
    return slices, labels

# -----------------------------------------------------------------------------
# FUNCÃO AUXILIAR QUE CONVERTE SLICES/LEGENDAS PARA ADICIONAR NO PARAGRAFO
# -----------------------------------------------------------------------------
def convert_itens_to_sequential_text(itens_values, itens_legends):
    text = ''
    for i in range(len(itens_values)):
        text = text + '{0}% ({1})'.format(itens_values[i], itens_legends[i].replace('\\',''))
        if i < len(itens_values) - 2:
            text = text + ', '
        elif i < len(itens_values) - 1:
            text = text + ' e '
            
    return text


# -----------------------------------------------------------------------------
# FUNÇÕES UTILIZADAS PARA ENVIAR EMAIL
# -----------------------------------------------------------------------------
def send_report_to_email(toaddr, filename):
    fromaddr = "YOUR_EMAIL"
    msg = MIMEMultipart()
     
    msg['From'] = fromaddr
    msg['To'] = toaddr
    msg['Subject'] = "Relatório Beta (Jogos para Console)"
    
    print('Enviando email para {0}'.format(msg['To']))
     
    body = "Prezado aluno, você está recebendo um email com o relatório beta anexado. Tenha um bom final de semana!"
     
    msg.attach(MIMEText(body, 'plain'))

    attachment = open(filename, "rb")
     
    part = MIMEBase('application', 'octet-stream')
    part.set_payload((attachment).read())
    encoders.encode_base64(part)
    part.add_header('Content-Disposition', "attachment; filename= %s" % filename)
     
    msg.attach(part)
     
    server = smtplib.SMTP('smtp.gmail.com', 587)
    server.starttls()
    server.login(fromaddr, "YOUR_PASSWORD")
    text = msg.as_string()
    server.sendmail(fromaddr, toaddr, text)
    server.quit()


# -----------------------------------------------------------------------------
# PRINCIPAL FUNÇÃO DO PROJETO (CRIAR O DOCUMENTO)
# -----------------------------------------------------------------------------
def make_documents():
    
    emails = {'NOME_DO_JOGO':'EMAIL_DO_ALUNO'}
    
    # lendo dados do excel
    data = pd.read_excel(r'beta evaluation.xlsx')
    
    games_full = data[COLUMN_GAME]
    ages = data[COLUMN_AGES]
    genders = data[COLUMN_GENDERS]
    graphics = data[COLUMN_GRAPHICS]
    hud = data[COLUMN_HUD]
    gameplays = data[COLUMN_GAMEPLAY]
    audios = data[COLUMN_AUDIO]
    stories = data[COLUMN_STORY]
    difficulties = data[COLUMN_DIFFICULTY]
    strongs = data[COLUMN_STRONG]
    weaks = data[COLUMN_WEAK]
    
    # Nome dos jogos
    games = list(set(games_full))   
    
    # Cria o documento para cada jogo
    for game in games:
        indexes = [i for i, x in enumerate(games_full) if x == game]
        
        print('Jogo {0} sendo avaliado...'.format(game))
        
        document = Document() 
        
        add_document_heading(document, REPORT_HEADING)  
        
        # Paragrafo 1
        add_document_paragraph(document, REPORT_PARAGRAPH_1)
        
        # Paragrafo 2
        male = np.count_nonzero(genders[indexes] == COLUMN_GENDERS_ITENS[0])
        female = np.count_nonzero(genders[indexes] == COLUMN_GENDERS_ITENS[1])
        
        ages_mean_std = str(int(round(np.mean(ages[indexes]),0)))
        
        stories_mean_std = str(round(np.mean(stories[indexes]),2)) + ' ± ' + str(round(np.std(stories[indexes]), 2))
        audios_mean_std = str(round(np.mean(audios[indexes]),2)) + ' ± ' + str(round(np.std(audios[indexes]), 2))
        gameplay_mean_std = str(round(np.mean(gameplays[indexes]),2)) + ' ± ' + str(round(np.std(gameplays[indexes]), 2))
        hud_mean_std = str(round(np.mean(hud[indexes]),2)) + ' ± ' + str(round(np.std(hud[indexes]), 2))
        graphics_mean_std = str(round(np.mean(graphics[indexes]),2)) + ' ± ' + str(round(np.std(graphics[indexes]), 2))
        
        add_document_paragraph(document, REPORT_PARAGRAPH_2.format(game, len(indexes), female, male, ages_mean_std, stories_mean_std, audios_mean_std, gameplay_mean_std, hud_mean_std, graphics_mean_std))
        
        #Figura 1
        plot_bar_err(game, [graphics[indexes], hud[indexes], gameplays[indexes], audios[indexes], stories[indexes]], [COLUMN_GRAPHICS, COLUMN_HUD, COLUMN_GAMEPLAY, COLUMN_AUDIO, COLUMN_STORY])
        add_document_figure(document, game + '_bar_err.png', REPORT_CAPTION_1, 9)
        
        
        # Pontos Fortes
        add_document_paragraph(document, REPORT_PARAGRAPH_4)
        add_document_list(document, strongs[indexes])
        add_document_paragraph(document, REPORT_PARAGRAPH_6)
        plot_word_cloud(game + '_strong_word_cloud.png', strongs[indexes])
        add_document_figure(document, game + '_strong_word_cloud.png', REPORT_CAPTION_3, 9)
        
        # Pontos Fracos
        add_document_paragraph(document, REPORT_PARAGRAPH_5)
        add_document_list(document, weaks[indexes])
        add_document_paragraph(document, REPORT_PARAGRAPH_7)
        plot_word_cloud(game + '_weak_word_cloud.png', weaks[indexes])
        add_document_figure(document, game + '_weak_word_cloud.png', REPORT_CAPTION_4, 9)
        
        # Paragrafo 3
        slices, legends = plot_pie_difficulty(game + '_difficulty.png', difficulties[indexes])
        diffilcuty_text = convert_itens_to_sequential_text(slices, legends)
        add_document_paragraph(document, REPORT_PARAGRAPH_3.format(diffilcuty_text))
        add_document_figure(document, game + '_difficulty.png', REPORT_CAPTION_2, 8)
        
        print('Salvando o documento {0}'.format(game + '.docx'))
        document.save(game + '.docx')
        
        
        send_report_to_email(emails[game], game + '.docx')
    
if __name__ == '__main__':
    make_documents()
    