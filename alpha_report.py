# -*- coding: utf-8 -*-
"""
Created on Tue Oct  9 15:33:26 2018

@author: Kleber de Oliveira Andrade
@email: pdjkleber@gmail.com
"""




# -----------------------------------------------------------------------------
# BIBLIOTECAS UTILIZADAS PARA GRÁFICOS E LISTAS
# -----------------------------------------------------------------------------
import matplotlib.pyplot as plt
import pandas as pd
import numpy as np
import numpy.core.defchararray as np_f


# -----------------------------------------------------------------------------
# BIBLIOTECAS PARA CRIAÇÃO DO DOCX (PYTHON-DOCX)
# -----------------------------------------------------------------------------
from docx import Document
from docx.shared import Cm
from docx.shared import Pt
from docx.shared import RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


# -----------------------------------------------------------------------------
# OUTRAS BIBLIOTECAS (POUCO UTILIZADAS NESTE PROJETO)
# -----------------------------------------------------------------------------
import os
from datetime import date
# limpa os arquivos da pasta conforme .bat
"""
    Informações do arquivo clear.bat
    del *.png
    del *.docx
"""
os.system('clear.bat')


# -----------------------------------------------------------------------------
# CONSTANTES UTILIZADAS PARA CRIAÇÃO DO RELATÓRIO (DOCX)
# -----------------------------------------------------------------------------
REPORT_HEADING = '3.1.\tRELATÓRIO DO JOGO VERSÃO ALFA'
REPORT_PARAGRAPH_1 = 'Uma pesquisa foi realizada no dia 29 de setembro de 2018 com os alunos do 6º semestre na disciplina de Jogos para Console a fim de avaliar a versão alfa de cada jogo desenvolvido como proposta de atividade prática do semestre. O questionário apresenta 5 questões de múltipla escolha e uma dissertativa.'
REPORT_PARAGRAPH_2 = 'O jogo \"{0}\", foi apresentado e avaliado pelo público presente composto por {1} colegas de sala. A Figura 1 apresenta um gráfico de barras com desvio padrão que avalia as seguintes informações: originalidade ({2}), diversão ({3}) e mercado ({4}).'
REPORT_PARAGRAPH_3 = 'Para complementar a análise do questionário, foi identificado o número de pessoas que gostaria de jogar este jogo e quanto pagaria por ele. A Figura 2.a, ressalta que, {0}% dos respondentes demonstraram interesse, no entanto {1}% dizem que não. A Figura 2.b apresenta quantos pagariam pelo jogo em intervalos pré-definidos: {2}.'
REPORT_PARAGRAPH_4 = 'Para concluir a análise do jogo alfa, os participantes fizeram os seguintes comentários (copiado na íntegra):'
REPORT_CAPTION_1 = 'Figura 1. Gráfico com as avaliações dos alunos.'
REPORT_CAPTION_2 = 'Figura 2. Quantidade de pessoas que demonstram interesse em jogar seu jogo e quanto pagariam por seu jogo.' 
REPORT_OWN_QUOTATION = 'Fonte: autoria própria ({0}).'.format(int(date.today().year))


# -----------------------------------------------------------------------------
# FUNÇÕES UTILIZADAS PARA CRIAÇÃO DO RELATÓRIO (DOCX)
# -----------------------------------------------------------------------------

def add_document_heading(document, heading):
    
    # Definição de um estilo para a seção do capítulo
    style = document.styles['Body Text 2']  
    font = style.font
    font.name = 'Arial'
    font.size = Pt(12)
    font.all_caps = True
    font.bold = True
    font.color.rgb = RGBColor(0x0, 0x0, 0x0)
    
    # Adiciona uma seção no documento
    document.add_paragraph(REPORT_HEADING, 'Body Text 2')  
    
    
def add_document_paragraph(document, text):
    
    # Definição de um estilo para o parágrafo
    style = document.styles['Body Text']    
    font = style.font
    font.name = 'Arial'
    font.size = Pt(12)
    
    # Adiciona um paragrafo no documento
    paragraph = document.add_paragraph(text, 'Body Text')
    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.first_line_indent = Cm(1.25)


def add_document_figure(document, figure, caption, width):
    
    # Definição de um estilo para a legenda da figura
    style = document.styles['Caption']    
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    font.bold = False
    font.color.rgb = RGBColor(0x0, 0x0, 0x0)
    
    # Adiciona legenda para figura
    document.add_paragraph()
    document.add_paragraph(caption, 'Caption')    
    
    # Adiciona uma figura no documento
    document.add_picture(figure, width=Cm(width))
    last_paragraph = document.paragraphs[-1] 
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Adiciona a fonte da figura (citação)
    document.add_paragraph(REPORT_OWN_QUOTATION, 'Caption')
    document.add_paragraph()
    
    
def add_document_list(document, itens):
    
    # Definição de um estilo para a lista de itens
    style = document.styles['List Bullet']    
    font = style.font
    font.name = 'Arial'
    font.size = Pt(12)
    
    # Adiciona os itens (textos) no documento
    for text in list(set(itens)):
        paragraph = document.add_paragraph(str(text).rstrip('\r').replace('\n', '; '), 'List Bullet')
        paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.first_line_indent = Cm(1.25)
    

# -----------------------------------------------------------------------------
# CONSTANTES DAS COLUNAS DO ARQUIVO EXCEL
# -----------------------------------------------------------------------------
COLUMN_GAME = 'Jogo avaliado'
COLUMN_ORIGINALITY = 'Originalidade'
COLUMN_FUN = 'Diversão'
COLUMN_MARKET = 'Mercado'
COLUMN_SUGGESTIONS = 'Críticas e Sugestões'
COLUMN_PLAY = 'Você jogaria este jogo?'
COLUMN_PLAY_ITENS = ['Sim', 'Não']
COLUMN_PAY = 'Quanto pagaria por esta ideia?'
COLUMN_PAY_ITENS = ['Menos R$ 10,00', 'Entre R$ 10,00 e R$ 20,00', 'Entre R$ 10,00 e R$ 20,00', 'Entre R$ 20,00 e R$ 30,00', 'Entre R$ 30,00 e R$ 50,00', 'Entre R$ 50,00 e R$ 75,00', 'Entre R$ 75,00 e R$ 100,00', 'Mais R$ 100,00']

# -----------------------------------------------------------------------------
# FUNÇÕES UTILIZADAS PARA CRIAÇÃO DOS GRÁFICOS
# -----------------------------------------------------------------------------
def plot_bar_err(filename, data):
    
    x = [COLUMN_ORIGINALITY, COLUMN_FUN, COLUMN_MARKET]
    x_pos = [i for i, _ in enumerate(x)]
    
    average = [np.mean(data[0]), np.mean(data[1]), np.mean(data[2])]
    variance = [np.std(data[0]), np.std(data[1]), np.std(data[2])]
    
    plt.bar(x_pos, average, color='steelblue', yerr=variance)
    
    plt.xticks(x_pos, x, fontsize=12)
    plt.yticks(list(range(11)), fontsize=12)
    
    plt.savefig(filename + '_bar.png', bbox_inches='tight', dpi=400)    
    plt.show()
    plt.close()
    

def plot_pie(filename, data1, data2):

    fig, axes = plt.subplots(nrows=1, ncols=2)
    
    # Gráfico de pizza (esquerda) correspondente a quantidade de jogadores com interesse em jogar
    length = len(data1) 
    py = round(np.count_nonzero(data1 == COLUMN_PAY_ITENS[0]) / length * 100, 0)
    pn = round(np.count_nonzero(data1 == COLUMN_PAY_ITENS[1]) / length * 100, 0)

    patches, texts, autotexts = axes[0].pie([py, pn], labels=COLUMN_PLAY_ITENS, autopct='%.1f%%', startangle=90, counterclock=False)
    axes[0].set_title(COLUMN_PLAY, fontsize=8)
    axes[0].axis('equal')
    
    plt.setp(autotexts, size=12, weight="bold")
    
    for t in texts:
        t.set_size('x-small')
    for t in autotexts:
        t.set_size('x-small')       
        

    # Gráfico de pizza (direita) correspondente a quanto pagariam pelo jogo
    labels, counts_elements = np.unique(data2, return_counts=True)
    slices = np.round(counts_elements / length * 100,1 )
    
    # Convertendo os labels para string e adicionando caracter de escape no $ (latex)
    labels_str = labels.astype(str)
    labels_str = np_f.replace(labels_str, '$', '\$')
    
    patches, texts, autotexts = axes[1].pie(slices, labels=labels_str, autopct='%.1f%%', startangle=90, counterclock=False)
    axes[1].set_title(COLUMN_PAY, fontsize=8)
    axes[1].axis('equal')
    
    plt.setp(autotexts, size=8, weight="bold") 

    for t in texts:
        t.set_size('x-small')
    for t in autotexts:
        t.set_size('x-small')
    
    plt.axis('equal')
    plt.subplots_adjust(wspace=0.75)
    plt.savefig(filename + '_pie.png', bbox_inches='tight', dpi=400)    
    plt.show()
    plt.close()

    return py, pn, slices, labels
    

# -----------------------------------------------------------------------------
# FUNCÃO AUXILIAR QUE CONVERTE SLICES/LEGENDAS PARA ADICIONAR NO PARAGRAFO
# -----------------------------------------------------------------------------
def convert_itens_to_sequential_text(itens_values, itens_legends):
    text = ''
    for i in range(len(itens_values)):
        text = text + '{0}% ({1})'.format(itens_values[i], itens_legends[i].replace('\\',''))
        if i < len(itens_values) - 2:
            text = text + ', '
        elif i < len(itens_values) - 1:
            text = text + ' e '
            
    return text


# -----------------------------------------------------------------------------
# PRINCIPAL FUNÇÃO DO PROJETO (CRIAR O DOCUMENTO)
# -----------------------------------------------------------------------------
def make_documents():
    
    # lendo dados do excel
    data = pd.read_excel(r'alpha evaluation.xlsx')
    games_full = data[COLUMN_GAME]
    originality = data[COLUMN_ORIGINALITY];
    fun = data[COLUMN_FUN];
    market = data[COLUMN_MARKET];
    you_would_play = data[COLUMN_PAY];
    suggestions = data[COLUMN_SUGGESTIONS]
    pay = data[COLUMN_PAY]
    
    # Nome dos jogos
    games = list(set(games_full))        
            
    # Cria o documento para cada jogo
    for game in games:
        indexes = [i for i, x in enumerate(games_full) if x == game]
        
        document = Document() 
        
        add_document_heading(document, REPORT_HEADING)  
        
        add_document_paragraph(document, REPORT_PARAGRAPH_1)
        
        orig_mean_std = str(round(np.mean(originality[indexes]),2)) + ' ± ' + str(round(np.std(originality[indexes]), 2))
        fun_mean_std = str(round(np.mean(fun[indexes]),2)) + ' ± ' + str(round(np.std(fun[indexes]), 2))
        mark_mean_std = str(round(np.mean(market[indexes]),2)) + ' ± ' + str(round(np.std(market[indexes]), 2))
        add_document_paragraph(document, REPORT_PARAGRAPH_2.format(game, len(indexes), orig_mean_std, fun_mean_std, mark_mean_std))
        
        plot_bar_err(game, [originality[indexes], fun[indexes], market[indexes]])
        
        add_document_figure(document, game + '_bar.png', REPORT_CAPTION_1, 12)
        
        py, pn, slices, legends = plot_pie(game, you_would_play[indexes], pay[indexes])
        
        text = convert_itens_to_sequential_text(slices, legends)
        add_document_paragraph(document, REPORT_PARAGRAPH_3.format(py, pn, text))
        
        add_document_figure(document, game + '_pie.png', REPORT_CAPTION_2, 15)   
        
        add_document_paragraph(document, REPORT_PARAGRAPH_4)
        
        add_document_list(document, suggestions[indexes])
        
        document.save(game + '.docx')
    
    
if __name__ == '__main__':
    make_documents()